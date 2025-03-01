import random
from docx import Document

def generate_word_search(words, grid_size=12):
    """
    Generates a word search puzzle with given Korean words.

    Args:
      words: A list of Korean words to include in the puzzle.
      grid_size: The size of the grid (default: 12).

    Returns:
      A tuple containing:
        - grid: A 2D list representing the word search grid.
        - word_locations: A dictionary mapping words to their starting coordinates and directions.
    """
    grid = [[' ' for _ in range(grid_size)] for _ in range(grid_size)]
    word_locations = {}

    for word in words:
        # Generate random starting coordinates and direction
        row, col = random.randint(0, grid_size - 1), random.randint(0, grid_size - 1)
        direction = random.choice(['horizontal', 'vertical', 'diagonal_down', 'diagonal_up'])

        # Adjust starting coordinates if necessary to fit the word in the grid
        if direction == 'horizontal':
            if col + len(word) > grid_size:
                col = grid_size - len(word)
        elif direction == 'vertical':
            if row + len(word) > grid_size:
                row = grid_size - len(word)
        elif direction == 'diagonal_down':
            if row + len(word) > grid_size or col + len(word) > grid_size:
                row, col = max(0, grid_size - len(word)), max(0, grid_size - len(word))
        elif direction == 'diagonal_up':
            if row - len(word) < 0 or col + len(word) > grid_size:
                row, col = min(grid_size - 1, len(word) - 1), max(0, grid_size - len(word))

        # Place the word in the grid
        if direction == 'horizontal':
            for i in range(len(word)):
                grid[row][col + i] = word[i]
        elif direction == 'vertical':
            for i in range(len(word)):
                grid[row + i][col] = word[i]
        # for diagonal durections if needed
        # elif direction == 'diagonal_down':
        #     for i in range(len(word)):
        #         grid[row + i][col + i] = word[i]
        # elif direction == 'diagonal_up':
        #     for i in range(len(word)):
        #         grid[row - i][col + i] = word[i]

    word_locations[word] = {'row': row, 'col': col, 'direction': direction}

    # Fill empty cells with random Korean syllables(i asked ChatGPT for most common syllables)
    korean_characters =  ["가", "개", "갸", "거", "게", "겨", "고", "구", "기", "나", "내", "냐", "너", "네", "녀", "노", "누", "니",
                          "다", "대", "댜", "더", "데", "뎌", "도", "두", "디", "라", "래", "랴", "러", "레", "려", "로", "루", "리",
                          "마", "매", "먀", "머", "메", "며", "모", "무", "미", "바", "배", "뱌", "버", "베", "벼", "보", "부", "비",
                          "사", "새", "샤", "서", "세", "셔", "소", "수", "시", "아", "애", "야", "어", "에", "여", "오", "우", "이",
                          "자", "재", "쟈", "저", "제", "져", "조", "주", "지", "차", "채", "챠", "처", "체", "쳐", "초", "추", "치",
                          "카", "캐", "캬", "커", "케", "켜", "코", "쿠", "키", "타", "태", "탸", "터", "테", "텨", "토", "투", "티",
                          "파", "패", "퍄", "퍼", "페", "펴", "포", "푸", "피", "하", "해", "햐", "허", "헤", "혀", "호", "후", "히",
                          "갖", "같", "개", "객", "갠", "걀", "걷", "걸", "검", "겁", "겉", "결", "겹", "경", "곁", "공", "과", "관",
                          "광", "괜", "괴", "굉", "교", "굳", "귀", "극", "근", "글", "깊", "꽂", "꽃", "꿈", "끝", "나", "난", "날",
                          "남", "납", "낫", "낮", "내", "냇", "냉", "너", "넉", "널", "넓", "넘", "넣", "녹", "논", "놀", "농", "높",
                          "눈", "눕", "느", "늑", "늦", "다", "단", "달", "닮", "담", "답", "당", "닿", "대", "댁", "댐", "덜", "덤",
                          "덥", "덧", "덩", "덫", "덴", "뎌", "도", "독", "돈", "돌", "동", "돼", "되", "된", "두", "둔", "둘", "둠",
                          "둡", "뒤", "뒷", "드", "득", "들", "등", "따", "딴", "딸", "땀", "땅", "때", "떠", "떡", "떨", "떻", "또",
                          "똑", "똥", "뚜", "뜯", "뜻", "라", "락", "란", "람", "랍", "랑", "래", "랜", "랭", "랴", "량", "러", "럭",
                          "런", "럴", "럼", "럽", "렇", "레", "렌", "려", "력", "련", "렬", "렴", "렵", "령", "례", "로", "록", "론",
                          "롤", "롬", "롱", "뢰", "료", "룡", "루", "룩", "룬", "룰", "룸", "룻", "뤼", "류", "률", "륨", "륭", "르",
                          "륵", "른", "를", "름", "릇", "리", "릭", "린", "림", "립", "링", "마", "막", "만", "많", "말", "망", "맞",
                          "매", "맥", "맴", "맵"
                          ]
    for i in range(grid_size):
        for j in range(grid_size):
            if grid[i][j] == ' ':
                grid[i][j] = random.choice(korean_characters)

    return grid, word_locations

#Add word-translation dictionary
def word_translate(korean_words,translations):
    """
    Creates dictionary with translation to each word.
    """
    word_translate = {}
    for korean_word, translation in zip(korean_words,translations):
        word_translate[korean_word] = translation
    return word_translate

def save_to_word(grid, word_translate, filename="word_search.docx"):
    """
    Saves the word search grid and translation dictionary to a Word document.

    Args:
      grid: The 2D list representing the word search grid.
      word_translate: dictionary with translations.
      filename: The name of the Word document to save (default: "word_search.docx").
    """

    doc = Document()

    # Add the word search grid to the document
    for row in grid:
        table = doc.add_table(rows=1, cols=len(row))
        table.style = 'Table Grid'
        row_cells = table.rows[0].cells
        for j, cell in enumerate(row_cells):
            cell.text = row[j]
    # Add word-translation dictionary to the document
    for word, transaltion in word_translate.items():
        doc.add_paragraph(f'{word} - {transaltion}')

    # Save the document
    doc.save(filename)



# Example usage
korean_words = ["김치", "비빔밥", "불고기", "떡볶이", "삼겹살", "라면", "갈비", "잡채", "된장찌개", "호떡"]
translations = ["Kimchi", "Bibimbap", "Bulgogi (grilled marinated beef)", 
                "Tteokbokki (spicy rice cakes)", "Samgyeopsal (grilled pork belly)", 
                "Ramen", "Galbi (Korean BBQ ribs)", "Japchae (stir-fried glass noodles)", 
                "Doenjang-jjigae (soybean paste stew)", "Hotteok (sweet Korean pancake)"]

grid, _ = generate_word_search(korean_words)
word_translate = word_translate(korean_words, translations)

save_to_word(grid, word_translate)
